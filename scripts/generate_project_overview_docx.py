from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/xlw/Documents/codex1/AI封面创意助手项目/ai-cover-creative-assistant")
SOURCE_MD = ROOT / "docs/product/AI封面创意助手_项目完整介绍_v0.1.md"
OUTPUT_DOCX = ROOT / "docs/product/AI封面创意助手_项目完整介绍_v0.1.docx"
BASE_FONT = "Hiragino Sans GB"


def set_run_font(run, font_name, size_pt, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        style.font.bold = False

    h1 = document.styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = document.styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = document.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(67, 67, 67)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    if "Quote Box" not in document.styles:
        style = document.styles.add_style("Quote Box", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Inches(0.15)
        style.paragraph_format.right_indent = Inches(0.15)


def add_title_block(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("AI封面创意助手\n项目完整介绍")
    set_run_font(run, BASE_FONT, 24, color=(0, 0, 0))

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run("版本：v0.1  |  文档用途：项目介绍 / 作品集材料 / 方案对齐  |  日期：2026-08-07")
    set_run_font(run, BASE_FONT, 10, color=(85, 85, 85))

    lead = document.add_paragraph(style="Quote Box")
    lead.paragraph_format.space_after = Pt(12)
    run = lead.add_run(
        "这是一份围绕 AI封面创意助手 的完整项目介绍文档，重点说明项目定位、目标用户、核心问题、产品方案、规则体系、工程架构与后续验证方向。"
    )
    set_run_font(run, BASE_FONT, 11)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_summary_table(document):
    document.add_paragraph("项目摘要", style="Heading 1")
    table = document.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    rows = [
        ("产品定位", "面向成长型内容创作者的垂直 AI 封面创意决策工具"),
        ("核心目标", "提升内容发布前的封面点击率与标题判断效率"),
        ("关键能力", "三方向卡输出、工作区深化、二轮修订、案例复盘"),
        ("方法论", "规则系统 + 案例库 + 大模型生成"),
        ("当前阶段", "主链路已成型，继续补规则与案例闭环，UI 优化后置"),
    ]
    for i, (left, right) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.9)
        for idx, text in enumerate((left, right)):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            set_run_font(run, BASE_FONT, 10.5, bold=(idx == 0))
            if idx == 0:
                shade_cell(cell, "F8F9FA")
    set_table_borders(table)
    document.add_paragraph()


def add_footer(document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AI封面创意助手项目完整介绍")
    set_run_font(run, BASE_FONT, 9, color=(85, 85, 85))


def add_markdown_content(document, markdown_text):
    lines = markdown_text.splitlines()
    first_title_skipped = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            if not first_title_skipped:
                first_title_skipped = True
                continue
            paragraph = document.add_paragraph(line[2:].strip(), style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            continue
        if line.startswith("## "):
            paragraph = document.add_paragraph(line[3:].strip(), style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            continue
        if line.startswith("### "):
            paragraph = document.add_paragraph(line[4:].strip(), style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(4)
            bullet = paragraph.add_run("• ")
            set_run_font(bullet, BASE_FONT, 11)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, BASE_FONT, 11)
            continue
        if line[:2].isdigit() and line[1:3] == ". ":
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(line)
            set_run_font(run, BASE_FONT, 11)
            continue
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(line)
        set_run_font(run, BASE_FONT, 11)


def build_docx():
    markdown_text = SOURCE_MD.read_text(encoding="utf-8")
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_title_block(document)
    add_summary_table(document)
    add_markdown_content(document, markdown_text)
    add_footer(document)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
